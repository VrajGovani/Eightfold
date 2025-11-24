"""
File parsing utilities for resume uploads
"""
import PyPDF2
import docx
import io
import logging
from typing import Optional

logger = logging.getLogger(__name__)

class FileParser:
    """Parse various file formats"""
    
    @staticmethod
    async def parse_pdf(file_content: bytes) -> str:
        """Extract text from PDF"""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error parsing PDF: {str(e)}")
            raise ValueError(f"Failed to parse PDF: {str(e)}")
    
    @staticmethod
    async def parse_docx(file_content: bytes) -> str:
        """Extract text from DOCX"""
        try:
            docx_file = io.BytesIO(file_content)
            doc = docx.Document(docx_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error parsing DOCX: {str(e)}")
            raise ValueError(f"Failed to parse DOCX: {str(e)}")
    
    @staticmethod
    async def parse_txt(file_content: bytes) -> str:
        """Extract text from TXT"""
        try:
            return file_content.decode('utf-8').strip()
        except Exception as e:
            logger.error(f"Error parsing TXT: {str(e)}")
            raise ValueError(f"Failed to parse TXT: {str(e)}")
    
    @staticmethod
    async def parse_file(filename: str, file_content: bytes) -> str:
        """Parse file based on extension"""
        extension = filename.lower().split('.')[-1]
        
        if extension == 'pdf':
            return await FileParser.parse_pdf(file_content)
        elif extension in ['docx', 'doc']:
            return await FileParser.parse_docx(file_content)
        elif extension == 'txt':
            return await FileParser.parse_txt(file_content)
        else:
            raise ValueError(f"Unsupported file format: {extension}")
